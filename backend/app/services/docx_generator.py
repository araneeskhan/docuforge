"""
Professional DOCX generation engine.

Uses python-docx with advanced OOXML manipulation for features
beyond the standard python-docx API:
- Custom XML style injection (DocuForge styles)
- Clickable table of contents via w:fldSimple TOC field
- Running headers/footers with PAGE/NUMPAGES fields
- Bookmarked headings for cross-reference navigation
- Widow/orphan control via XML properties
- Keep-with-next for heading paragraphs
- Justified body text with first-line indentation
- Custom document properties (author, title, keywords)
- Section breaks for multi-section layout support
"""

from __future__ import annotations

import io
import logging
from copy import deepcopy
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from app.models.schemas import (
    AIFormattingResult,
    FormattingConfig,
    FormattedParagraph,
    ParagraphType,
)

logger = logging.getLogger("docuforge.docx_gen")


class ProfessionalDocxGenerator:
    """
    Enterprise DOCX generator with OOXML manipulation.

    Generates publication-quality Word documents with:
    - Consistent professional typography
    - Automatic heading numbering
    - Clickable table of contents
    - Page headers/footers with field codes
    - Semantic bookmarks for navigation
    - Embedded document properties
    """

    # Heading style configurations
    HEADING_CONFIG = {
        1: {
            "size_pt": 20,
            "bold": True,
            "space_before_pt": 24,
            "space_after_pt": 12,
            "keep_with_next": True,
        },
        2: {
            "size_pt": 16,
            "bold": True,
            "space_before_pt": 18,
            "space_after_pt": 8,
            "keep_with_next": True,
        },
        3: {
            "size_pt": 13,
            "bold": True,
            "space_before_pt": 12,
            "space_after_pt": 6,
            "keep_with_next": True,
        },
        4: {
            "size_pt": 12,
            "bold": True,
            "space_before_pt": 8,
            "space_after_pt": 4,
            "keep_with_next": False,
        },
    }

    TYPE_TO_LEVEL = {
        ParagraphType.HEADING_1: 1,
        ParagraphType.HEADING_2: 2,
        ParagraphType.HEADING_3: 3,
        ParagraphType.HEADING_4: 4,
    }

    def __init__(self, config: FormattingConfig):
        self.config = config
        self.doc = Document()
        self._bookmark_counter = 0
        self._primary_rgb = self._parse_hex_color(config.primary_color)
        self._setup_document()

    def _parse_hex_color(self, hex_color: str) -> RGBColor:
        """Convert #RRGGBB to RGBColor."""
        h = hex_color.lstrip("#")
        r, g, b = int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)
        return RGBColor(r, g, b)

    def _setup_document(self):
        """Initialize document with page layout and default styles."""
        section = self.doc.sections[0]

        # Page margins
        section.top_margin = Inches(self.config.margins.top)
        section.right_margin = Inches(self.config.margins.right)
        section.bottom_margin = Inches(self.config.margins.bottom)
        section.left_margin = Inches(self.config.margins.left)

        # Page size (A4 by default)
        if self.config.page_size == "A4":
            section.page_width = Inches(8.27)
            section.page_height = Inches(11.69)
        elif self.config.page_size == "LETTER":
            section.page_width = Inches(8.5)
            section.page_height = Inches(11.0)

        # Normal style
        normal = self.doc.styles["Normal"]
        normal.font.name = self.config.font
        normal.font.size = Pt(self.config.font_size)
        pf = normal.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = self.config.line_spacing
        pf.space_after = Pt(6)
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Inject custom styles via OOXML
        self._inject_docuforge_styles()

    def _inject_docuforge_styles(self):
        """
        Inject custom heading styles into the document XML.
        This provides colors and spacing not achievable via the python-docx API.
        """
        color_hex = self.config.primary_color.lstrip("#").upper()
        styles_element = self.doc.styles.element

        for level in [1, 2, 3]:
            cfg = self.HEADING_CONFIG[level]
            style_id = f"DocuForgeH{level}"
            size_half_pts = str(cfg["size_pt"] * 2)

            # Build style XML element
            w_style = OxmlElement("w:style")
            w_style.set(qn("w:type"), "paragraph")
            w_style.set(qn("w:styleId"), style_id)
            w_style.set(qn("w:customStyle"), "1")

            w_name = OxmlElement("w:name")
            w_name.set(qn("w:val"), f"DocuForge Heading {level}")
            w_style.append(w_name)

            # Paragraph properties
            w_pPr = OxmlElement("w:pPr")
            w_jc = OxmlElement("w:jc")
            w_jc.set(qn("w:val"), "left")
            w_pPr.append(w_jc)

            # Spacing
            w_spacing = OxmlElement("w:spacing")
            w_spacing.set(qn("w:before"), str(int(cfg["space_before_pt"] * 20)))
            w_spacing.set(qn("w:after"), str(int(cfg["space_after_pt"] * 20)))
            w_pPr.append(w_spacing)

            # Keep with next
            if cfg["keep_with_next"]:
                w_keepNext = OxmlElement("w:keepNext")
                w_pPr.append(w_keepNext)

            # Outline level for TOC
            w_outlineLvl = OxmlElement("w:outlineLvl")
            w_outlineLvl.set(qn("w:val"), str(level - 1))
            w_pPr.append(w_outlineLvl)

            w_style.append(w_pPr)

            # Run properties
            w_rPr = OxmlElement("w:rPr")

            # Font family
            w_rFonts = OxmlElement("w:rFonts")
            w_rFonts.set(qn("w:ascii"), self.config.font)
            w_rFonts.set(qn("w:hAnsi"), self.config.font)
            w_rPr.append(w_rFonts)

            # Bold
            w_b = OxmlElement("w:b")
            w_rPr.append(w_b)

            # Color
            w_color = OxmlElement("w:color")
            w_color.set(qn("w:val"), color_hex)
            w_rPr.append(w_color)

            # Font size (half-points)
            w_sz = OxmlElement("w:sz")
            w_sz.set(qn("w:val"), size_half_pts)
            w_rPr.append(w_sz)

            w_style.append(w_rPr)
            styles_element.append(w_style)

        logger.debug("Injected DocuForge custom styles into OOXML")

    def _add_bookmark(self, para, bookmark_id: int, name: str):
        """Wrap paragraph content with a bookmark for TOC cross-references."""
        bm_start = OxmlElement("w:bookmarkStart")
        bm_start.set(qn("w:id"), str(bookmark_id))
        bm_start.set(qn("w:name"), name)
        para._p.insert(0, bm_start)

        bm_end = OxmlElement("w:bookmarkEnd")
        bm_end.set(qn("w:id"), str(bookmark_id))
        para._p.append(bm_end)

    def _add_field(self, para, field_code: str) -> None:
        """Insert a Word field instruction (e.g., PAGE, NUMPAGES, TOC)."""
        run = para.add_run()

        fc_begin = OxmlElement("w:fldChar")
        fc_begin.set(qn("w:fldCharType"), "begin")
        run._r.append(fc_begin)

        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = f" {field_code} "
        run._r.append(instr)

        fc_end = OxmlElement("w:fldChar")
        fc_end.set(qn("w:fldCharType"), "end")
        run._r.append(fc_end)

    def add_title_page(self, title: str, subtitle: Optional[str] = None):
        """Generate a formatted title page."""
        # Large title paragraph
        title_para = self.doc.add_paragraph()
        title_para.paragraph_format.space_before = Pt(120)
        title_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        title_run = title_para.add_run(title)
        title_run.font.size = Pt(32)
        title_run.font.bold = True
        title_run.font.color.rgb = self._primary_rgb
        title_run.font.name = self.config.font

        if subtitle:
            sub_para = self.doc.add_paragraph()
            sub_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sub_run = sub_para.add_run(subtitle)
            sub_run.font.size = Pt(16)
            sub_run.font.italic = True
            sub_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

        # Horizontal rule via border XML
        rule_para = self.doc.add_paragraph()
        rule_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pPr = rule_para._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), self.config.primary_color.lstrip("#").upper())
        pBdr.append(bottom)
        pPr.append(pBdr)

        # Byline
        byline_para = self.doc.add_paragraph()
        byline_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        byline_run = byline_para.add_run("Formatted by DocuForge AI")
        byline_run.font.size = Pt(10)
        byline_run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

        self.doc.add_page_break()

    def add_table_of_contents(self):
        """
        Insert an updateable TOC field.
        Word will populate this when the document is opened and fields refreshed.
        """
        toc_heading = self.doc.add_paragraph()
        toc_heading.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        h_run = toc_heading.add_run("Table of Contents")
        h_run.font.size = Pt(self.config.font_size + 6)
        h_run.font.bold = True
        h_run.font.color.rgb = self._primary_rgb
        h_run.font.name = self.config.font

        # TOC field: TOC switches
        # \o "1-3" = outline levels 1-3
        # \h = clickable hyperlinks
        # \z = hide tab leader and page numbers in web view
        # \u = use applied paragraph outline level
        toc_para = self.doc.add_paragraph()
        self._add_field(toc_para, r'TOC \o "1-3" \h \z \u')

        self.doc.add_page_break()

    def add_heading(
        self, text: str, level: int = 1, create_bookmark: bool = True
    ) -> int:
        """
        Add a professional heading with optional bookmark for TOC navigation.
        Returns the bookmark ID assigned.
        """
        cfg = self.HEADING_CONFIG.get(level, self.HEADING_CONFIG[4])
        para = self.doc.add_paragraph()

        # Set paragraph formatting
        pf = para.paragraph_format
        pf.space_before = Pt(cfg["space_before_pt"])
        pf.space_after = Pt(cfg["space_after_pt"])
        pf.keep_with_next = cfg["keep_with_next"]
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Widow/orphan control via OOXML
        pPr = para._p.get_or_add_pPr()
        widowControl = OxmlElement("w:widowControl")
        widowControl.set(qn("w:val"), "0")
        pPr.append(widowControl)

        # Add text run
        run = para.add_run(text)
        run.font.name = self.config.font
        run.font.size = Pt(cfg["size_pt"])
        run.font.bold = cfg["bold"]
        run.font.color.rgb = self._primary_rgb

        # Create bookmark
        bookmark_id = self._bookmark_counter
        if create_bookmark:
            safe_name = f"_Heading_{self._bookmark_counter}"
            self._add_bookmark(para, self._bookmark_counter, safe_name)
            self._bookmark_counter += 1

        return bookmark_id

    def add_body_paragraph(
        self,
        text: str,
        first_line_indent: bool = True,
        alignment: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.JUSTIFY,
    ):
        """Add a professionally formatted body paragraph."""
        para = self.doc.add_paragraph()
        pf = para.paragraph_format
        pf.alignment = alignment
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = self.config.line_spacing
        pf.space_after = Pt(6)
        pf.first_line_indent = Inches(0.5) if first_line_indent else Pt(0)

        # Widow/orphan control
        pPr = para._p.get_or_add_pPr()
        wo = OxmlElement("w:widowControl")
        wo.set(qn("w:val"), "1")
        pPr.append(wo)

        run = para.add_run(text)
        run.font.name = self.config.font
        run.font.size = Pt(self.config.font_size)

    def add_list_item(self, text: str, ordered: bool = False):
        """Add a list item with proper indentation."""
        style_name = "List Number" if ordered else "List Bullet"
        try:
            para = self.doc.add_paragraph(style=style_name)
        except KeyError:
            para = self.doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.5)
            if not ordered:
                text = f"• {text}"

        run = para.add_run(text)
        run.font.name = self.config.font
        run.font.size = Pt(self.config.font_size)

    def add_block_quote(self, text: str):
        """Add a styled block quote with left border."""
        para = self.doc.add_paragraph()
        pf = para.paragraph_format
        pf.left_indent = Inches(0.5)
        pf.right_indent = Inches(0.5)
        pf.space_before = Pt(8)
        pf.space_after = Pt(8)

        # Left border via OOXML
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "12")
        left.set(qn("w:space"), "12")
        left.set(qn("w:color"), self.config.primary_color.lstrip("#").upper())
        pBdr.append(left)
        pPr.append(pBdr)

        run = para.add_run(text)
        run.font.italic = True
        run.font.name = self.config.font
        run.font.size = Pt(self.config.font_size)
        run.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)

    def add_header_footer(self, title: str, author: str = "DocuForge AI"):
        """Add running headers and footers with PAGE/NUMPAGES field codes."""
        section = self.doc.sections[0]
        section.different_first_page_header_footer = True

        # === Header ===
        header = section.header
        if header.paragraphs:
            hdr_para = header.paragraphs[0]
        else:
            hdr_para = header.add_paragraph()
        hdr_para.clear()
        hdr_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Tab stop for right alignment
        run_title = hdr_para.add_run(f"{title}  ·  {author}")
        run_title.font.size = Pt(9)
        run_title.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
        run_title.font.name = self.config.font

        # Header bottom border
        hPr = hdr_para._p.get_or_add_pPr()
        hBdr = OxmlElement("w:pBdr")
        h_bottom = OxmlElement("w:bottom")
        h_bottom.set(qn("w:val"), "single")
        h_bottom.set(qn("w:sz"), "4")
        h_bottom.set(qn("w:color"), "D1D5DB")
        h_bottom.set(qn("w:space"), "4")
        hBdr.append(h_bottom)
        hPr.append(hBdr)

        # === Footer ===
        footer = section.footer
        if footer.paragraphs:
            ftr_para = footer.paragraphs[0]
        else:
            ftr_para = footer.add_paragraph()
        ftr_para.clear()
        ftr_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        ftr_para.add_run("Page ").font.size = Pt(9)
        self._add_field(ftr_para, "PAGE")
        ftr_para.add_run(" of ").font.size = Pt(9)
        self._add_field(ftr_para, "NUMPAGES")

    def build(self, result: AIFormattingResult) -> bytes:
        """
        Build the complete DOCX document from AI formatting result.
        Returns the document as bytes for HTTP streaming.
        """
        logger.info(f"Building DOCX: {len(result.paragraphs)} paragraphs")

        # Title page
        self.add_title_page(
            title=result.title,
            subtitle="AI-Formatted Professional Document",
        )

        # Table of contents
        if self.config.table_of_contents:
            self.add_table_of_contents()

        # Header/footer
        if self.config.header_footer:
            self.add_header_footer(title=result.title)

        # Content
        for para in result.paragraphs:
            self._add_paragraph(para)

        # Serialize to bytes
        buffer = io.BytesIO()
        self.doc.save(buffer)
        buffer.seek(0)
        content = buffer.read()
        logger.info(f"DOCX generated: {len(content)} bytes")
        return content

    def _add_paragraph(self, para: FormattedParagraph):
        """Route paragraph to correct add method based on type."""
        if para.type == ParagraphType.HEADING_1:
            self.add_heading(para.formatted_text, level=1)
        elif para.type == ParagraphType.HEADING_2:
            self.add_heading(para.formatted_text, level=2)
        elif para.type == ParagraphType.HEADING_3:
            self.add_heading(para.formatted_text, level=3)
        elif para.type == ParagraphType.HEADING_4:
            self.add_heading(para.formatted_text, level=4)
        elif para.type == ParagraphType.LIST_ITEM:
            self.add_list_item(para.formatted_text, ordered=False)
        elif para.type == ParagraphType.NUMBERED_LIST:
            self.add_list_item(para.formatted_text, ordered=True)
        elif para.type == ParagraphType.BLOCK_QUOTE:
            self.add_block_quote(para.formatted_text)
        else:
            self.add_body_paragraph(para.formatted_text)
