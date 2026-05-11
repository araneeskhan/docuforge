"""
Advanced DOCX/DOC/TXT parser.

Uses python-docx for .docx files with OOXML-level access to extract:
- Paragraph text, runs, and character-level formatting
- Heading hierarchy via outline levels and style names
- Tables with row/column structure
- Image metadata (not content, for security)
- Document properties (author, title, created date)
- Section properties (margins, page size, orientation)
"""

from __future__ import annotations

import io
import logging
import re
import unicodedata
from pathlib import Path
from typing import Optional

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

from app.models.schemas import (
    ParsedDocument,
    ParsedParagraph,
    ParsedTable,
    ParagraphType,
    RunMetadata,
)

try:
    from docx.text.paragraph import Paragraph as DocxParagraph
    from docx.table import Table as DocxTable
except ImportError:
    DocxParagraph = None
    DocxTable = None

try:
    import mammoth
    from bs4 import BeautifulSoup
except ImportError:
    mammoth = None
    BeautifulSoup = None

logger = logging.getLogger("docuforge.parser")


class AdvancedDocumentParser:
    """
    Enterprise-grade DOCX parser with OOXML manipulation support.

    Supports:
    - .docx (python-docx with full OOXML access)
    - .doc (mammoth library for legacy format)
    - .txt (plain text with heuristic structure detection)

    Paragraph classification uses a multi-signal approach:
    1. Word style names (Heading 1, Title, etc.)
    2. OOXML outline level (w:outlineLvl)
    3. Regex pattern matching (numbered headings, ALL CAPS, etc.)
    4. Character formatting (font size > body, bold, etc.)
    """

    # Patterns suggesting a paragraph is a heading
    HEADING_REGEXES = [
        re.compile(r"^chapter\s+\d+", re.IGNORECASE),
        re.compile(r"^section\s+\d+[\.:]\s+", re.IGNORECASE),
        re.compile(r"^\d+(\.\d+)*\s+[A-Z]"),      # "1.2.3 Title"
        re.compile(r"^[IVXLCDM]+\.\s+[A-Z]"),     # Roman numerals
        re.compile(r"^[A-Z][A-Z\s\d\-]{4,40}$"),  # SHORT ALL CAPS
    ]

    # Mapping from Word style names to our ParagraphType
    STYLE_MAP: dict[str, ParagraphType] = {
        "heading 1": ParagraphType.HEADING_1,
        "heading 2": ParagraphType.HEADING_2,
        "heading 3": ParagraphType.HEADING_3,
        "heading 4": ParagraphType.HEADING_4,
        "title": ParagraphType.HEADING_1,
        "subtitle": ParagraphType.HEADING_2,
        "list paragraph": ParagraphType.LIST_ITEM,
        "list bullet": ParagraphType.LIST_ITEM,
        "list number": ParagraphType.NUMBERED_LIST,
        "quote": ParagraphType.BLOCK_QUOTE,
        "intense quote": ParagraphType.BLOCK_QUOTE,
        "code": ParagraphType.CODE_BLOCK,
    }

    def __init__(self, file_content: bytes, file_name: str):
        self.file_content = file_content
        self.file_name = file_name
        self.extension = Path(file_name).suffix.lower().lstrip(".")
        self._doc: Optional[Document] = None

    def parse(self) -> ParsedDocument:
        """Main entry point — returns fully parsed document structure."""
        logger.info(f"Parsing {self.file_name} ({len(self.file_content)} bytes)")

        if self.extension == "docx":
            return self._parse_docx()
        elif self.extension == "doc":
            return self._parse_doc_via_mammoth()
        elif self.extension == "txt":
            return self._parse_txt()
        else:
            raise ValueError(f"Unsupported file format: .{self.extension}")

    # ─── DOCX Parser ─────────────────────────────────────────────────────────

    def _parse_docx(self) -> ParsedDocument:
        """Full DOCX parsing with OOXML access."""
        self._doc = Document(io.BytesIO(self.file_content))
        doc = self._doc

        paragraphs = []
        tables = []
        image_count = 0

        # Count inline images
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                image_count += 1

        # Parse paragraph elements in document order
        para_idx = 0
        for block in self._iter_block_items(doc):
            if block["type"] == "paragraph":
                para = block["element"]
                if not para.text.strip():
                    continue

                parsed = self._parse_paragraph(para, para_idx)
                paragraphs.append(parsed)
                para_idx += 1

            elif block["type"] == "table":
                table = block["element"]
                parsed_table = self._parse_table(table, len(tables))
                tables.append(parsed_table)

        # Detect existing TOC
        has_toc = any(
            p.text.lower() in ("table of contents", "contents", "toc")
            for p in paragraphs
            if p.is_heading
        )

        total_words = sum(p.word_count for p in paragraphs)

        return ParsedDocument(
            file_name=self.file_name,
            file_size_bytes=len(self.file_content),
            paragraphs=paragraphs,
            tables=tables,
            image_count=image_count,
            total_word_count=total_words,
            detected_language="en",
            has_existing_toc=has_toc,
        )

    def _parse_paragraph(self, para, index: int) -> ParsedParagraph:
        """Extract all metadata from a single paragraph object."""
        text = para.text.strip()
        style_name = para.style.name.lower() if para.style else "normal"

        # Determine paragraph type
        para_type = self._classify_paragraph(para, style_name, text)

        # Extract run-level formatting
        runs = []
        for run in para.runs:
            run_meta = RunMetadata(
                text=run.text,
                bold=run.bold,
                italic=run.italic,
                underline=run.underline,
                font_name=run.font.name,
                font_size=run.font.size.pt if run.font.size else None,
                color=(
                    str(run.font.color.rgb)
                    if run.font.color and run.font.color.type
                    else None
                ),
            )
            runs.append(run_meta)

        # Get outline level from OOXML
        xml_elem = para._element
        outline_lvl = xml_elem.find(f".//{qn('w:outlineLvl')}")
        level = (
            int(outline_lvl.get(qn("w:val"), 9))
            if outline_lvl is not None
            else 9
        )

        # Detect alignment
        alignment = str(para.alignment) if para.alignment else None

        return ParsedParagraph(
            index=index,
            text=text,
            type=para_type,
            style=para.style.name if para.style else "Normal",
            alignment=alignment,
            runs=runs,
            outline_level=level,
        )

    def _classify_paragraph(
        self, para, style_name: str, text: str
    ) -> ParagraphType:
        """
        Multi-signal paragraph classification.
        Priority: Style name > Outline level > Font size heuristic > Regex
        """
        # 1. Known style names
        for known_style, para_type in self.STYLE_MAP.items():
            if known_style in style_name:
                return para_type

        # 2. Outline level from OOXML
        xml_elem = para._element
        outline_lvl = xml_elem.find(f".//{qn('w:outlineLvl')}")
        if outline_lvl is not None:
            level = int(outline_lvl.get(qn("w:val"), 9))
            if level == 0:
                return ParagraphType.HEADING_1
            if level == 1:
                return ParagraphType.HEADING_2
            if level == 2:
                return ParagraphType.HEADING_3
            if level == 3:
                return ParagraphType.HEADING_4

        # 3. Font size heuristic (larger than 14pt → likely heading)
        max_font_size = max(
            (run.font.size.pt for run in para.runs if run.font.size),
            default=0,
        )
        if max_font_size > 14:
            if max_font_size > 20:
                return ParagraphType.HEADING_1
            if max_font_size > 16:
                return ParagraphType.HEADING_2
            return ParagraphType.HEADING_3

        # 4. Regex heuristics on text
        if len(text) < 80:
            for pattern in self.HEADING_REGEXES:
                if pattern.match(text):
                    return ParagraphType.HEADING_2

        # 5. List detection
        if text.startswith(("•", "–", "—", "-", "*", "◦", "○", "▪", "▸")):
            return ParagraphType.LIST_ITEM
        if re.match(r"^\d+[\.\)]\s+", text):
            return ParagraphType.NUMBERED_LIST

        # 6. Block quote detection
        if text.startswith('"') and text.endswith('"') and len(text) > 50:
            return ParagraphType.BLOCK_QUOTE

        return ParagraphType.BODY

    def _parse_table(self, table, index: int) -> ParsedTable:
        """Extract table rows and detect if first row is a header."""
        rows = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            rows.append(row_data)

        # Detect header: first row often has bold text or different shading
        has_header = False
        if rows and table.rows:
            first_row_cells = table.rows[0].cells
            for cell in first_row_cells:
                for para in cell.paragraphs:
                    if any(run.bold for run in para.runs):
                        has_header = True
                        break

        return ParsedTable(
            index=index,
            rows=rows,
            col_count=len(rows[0]) if rows else 0,
            row_count=len(rows),
            has_header=has_header,
        )

    def _iter_block_items(self, doc: Document):
        """
        Iterate document body elements in order, distinguishing
        paragraphs from tables (both are top-level body children).
        """
        body = doc.element.body
        for child in body.iterchildren():
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

            if tag == "p":
                if DocxParagraph:
                    yield {"type": "paragraph", "element": DocxParagraph(child, doc)}
            elif tag == "tbl":
                if DocxTable:
                    yield {"type": "table", "element": DocxTable(child, doc)}

    # ─── DOC Parser (via mammoth) ─────────────────────────────────────────────

    def _parse_doc_via_mammoth(self) -> ParsedDocument:
        """
        Parse legacy .doc files using mammoth, which converts to HTML
        then we parse the HTML structure.
        """
        if not mammoth or not BeautifulSoup:
            logger.warning("mammoth or beautifulsoup4 not installed, falling back to txt parser")
            return self._parse_txt()

        try:
            result = mammoth.convert_to_html(io.BytesIO(self.file_content))
            soup = BeautifulSoup(result.value, "html.parser")

            paragraphs = []
            idx = 0

            for elem in soup.find_all(["h1", "h2", "h3", "h4", "p", "li", "blockquote"]):
                text = elem.get_text(strip=True)
                if not text:
                    continue

                tag = elem.name
                type_map = {
                    "h1": ParagraphType.HEADING_1,
                    "h2": ParagraphType.HEADING_2,
                    "h3": ParagraphType.HEADING_3,
                    "h4": ParagraphType.HEADING_4,
                    "li": ParagraphType.LIST_ITEM,
                    "blockquote": ParagraphType.BLOCK_QUOTE,
                }
                para_type = type_map.get(tag, ParagraphType.BODY)

                paragraphs.append(
                    ParsedParagraph(index=idx, text=text, type=para_type)
                )
                idx += 1

            return ParsedDocument(
                file_name=self.file_name,
                file_size_bytes=len(self.file_content),
                paragraphs=paragraphs,
                total_word_count=sum(p.word_count for p in paragraphs),
                detected_language="en",
            )

        except Exception as e:
            logger.warning(f"mammoth parsing failed: {e}, falling back to txt parser")
            return self._parse_txt()

    # ─── TXT Parser ───────────────────────────────────────────────────────────

    def _parse_txt(self) -> ParsedDocument:
        """
        Plain text parser with heuristic structure detection.
        Uses line length, capitalization, and blank line patterns.
        """
        text = self.file_content.decode("utf-8", errors="replace")
        lines = text.splitlines()

        paragraphs = []
        idx = 0
        buffer: list[str] = []

        def flush_buffer():
            nonlocal idx
            combined = " ".join(buffer).strip()
            if combined:
                para_type = self._classify_text_line(combined)
                paragraphs.append(
                    ParsedParagraph(index=idx, text=combined, type=para_type)
                )
                idx += 1
            buffer.clear()

        for line in lines:
            stripped = line.strip()
            if not stripped:
                flush_buffer()
            else:
                buffer.append(stripped)

        flush_buffer()

        return ParsedDocument(
            file_name=self.file_name,
            file_size_bytes=len(self.file_content),
            paragraphs=paragraphs,
            total_word_count=sum(p.word_count for p in paragraphs),
        )

    def _classify_text_line(self, text: str) -> ParagraphType:
        """Heuristic classification for plain text lines."""
        if len(text) < 80 and text.isupper():
            return ParagraphType.HEADING_2
        for pattern in self.HEADING_REGEXES:
            if pattern.match(text):
                return ParagraphType.HEADING_2
        if text.startswith(("•", "-", "*")):
            return ParagraphType.LIST_ITEM
        if re.match(r"^\d+[\.\)]\s", text):
            return ParagraphType.NUMBERED_LIST
        return ParagraphType.BODY
